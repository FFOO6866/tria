"""
Document Processing for RAG Knowledge Base
===========================================

Extracts text from .docx files and chunks intelligently for embedding generation.

Features:
- Real .docx extraction using python-docx (NO MOCKS)
- Intelligent chunking with token limits and overlap
- Preserves document structure and metadata
- Production-ready error handling

Chunking Strategy:
- 500 tokens per chunk (optimal for embeddings)
- 100 token overlap (preserves context across boundaries)
- Uses tiktoken for accurate token counting
"""

import os
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
import tiktoken


# Token limits for chunking
DEFAULT_CHUNK_SIZE = 500  # tokens
DEFAULT_CHUNK_OVERLAP = 100  # tokens
ENCODING_MODEL = "cl100k_base"  # Same encoding as text-embedding-3-small


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract all text from a .docx file

    Args:
        file_path: Path to .docx file

    Returns:
        Complete text content from the document

    Raises:
        FileNotFoundError: If file doesn't exist
        RuntimeError: If document parsing fails
    """
    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    if not file_path.suffix.lower() == '.docx':
        raise ValueError(f"Not a .docx file: {file_path}")

    try:
        doc = Document(str(file_path))

        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only include non-empty paragraphs
                paragraphs.append(text)

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        # Join with double newlines to preserve document structure
        full_text = '\n\n'.join(paragraphs)

        if not full_text.strip():
            raise RuntimeError(f"Document appears to be empty: {file_path}")

        return full_text

    except Exception as e:
        raise RuntimeError(
            f"Failed to extract text from {file_path.name}. "
            f"Error: {str(e)}"
        ) from e


def count_tokens(text: str, encoding_name: str = ENCODING_MODEL) -> int:
    """
    Count tokens in text using tiktoken

    Args:
        text: Text to count tokens for
        encoding_name: Tokenizer encoding name

    Returns:
        Number of tokens
    """
    encoding = tiktoken.get_encoding(encoding_name)
    return len(encoding.encode(text))


def chunk_text(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    encoding_name: str = ENCODING_MODEL
) -> List[str]:
    """
    Split text into chunks with token limits and overlap

    Args:
        text: Full text to chunk
        chunk_size: Maximum tokens per chunk (default: 500)
        chunk_overlap: Overlap tokens between chunks (default: 100)
        encoding_name: Tokenizer encoding name

    Returns:
        List of text chunks

    Raises:
        ValueError: If chunk_size <= chunk_overlap
    """
    if chunk_size <= chunk_overlap:
        raise ValueError(f"chunk_size ({chunk_size}) must be greater than chunk_overlap ({chunk_overlap})")

    encoding = tiktoken.get_encoding(encoding_name)

    # Tokenize the full text
    tokens = encoding.encode(text)

    # If text is shorter than chunk size, return as single chunk
    if len(tokens) <= chunk_size:
        return [text]

    chunks = []
    start_idx = 0

    while start_idx < len(tokens):
        # Get chunk tokens
        end_idx = start_idx + chunk_size
        chunk_tokens = tokens[start_idx:end_idx]

        # Decode back to text
        chunk_text = encoding.decode(chunk_tokens)
        chunks.append(chunk_text)

        # Move start index forward (chunk_size - overlap)
        start_idx += (chunk_size - chunk_overlap)

        # Prevent infinite loop
        if start_idx <= start_idx - (chunk_size - chunk_overlap):
            break

    return chunks


def process_document(
    file_path: Path,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
) -> List[Dict[str, str]]:
    """
    Process a .docx document into chunks with metadata

    Args:
        file_path: Path to .docx file
        chunk_size: Maximum tokens per chunk
        chunk_overlap: Overlap tokens between chunks

    Returns:
        List of dictionaries with chunk text and metadata:
        [
            {
                'text': 'chunk text...',
                'source': 'filename.docx',
                'chunk_index': 0,
                'total_chunks': 5
            },
            ...
        ]

    Raises:
        FileNotFoundError: If file doesn't exist
        RuntimeError: If processing fails
    """
    # Extract full text
    full_text = extract_text_from_docx(file_path)

    # Chunk the text
    chunks = chunk_text(full_text, chunk_size, chunk_overlap)

    # Build chunk dictionaries with metadata
    processed_chunks = []
    total_chunks = len(chunks)

    for idx, text_chunk in enumerate(chunks):
        processed_chunks.append({
            'text': text_chunk,
            'source': file_path.name,
            'chunk_index': idx,
            'total_chunks': total_chunks
        })

    return processed_chunks


def process_directory(
    directory: Path,
    file_pattern: str = "*.docx",
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
) -> Dict[str, List[Dict[str, str]]]:
    """
    Process all .docx files in a directory

    Args:
        directory: Directory containing .docx files
        file_pattern: Glob pattern for files (default: *.docx)
        chunk_size: Maximum tokens per chunk
        chunk_overlap: Overlap tokens between chunks

    Returns:
        Dictionary mapping filenames to their processed chunks

    Raises:
        FileNotFoundError: If directory doesn't exist
    """
    if not directory.exists():
        raise FileNotFoundError(f"Directory not found: {directory}")

    results = {}

    # Find all matching files
    docx_files = list(directory.glob(file_pattern))

    if len(docx_files) == 0:
        raise FileNotFoundError(f"No .docx files found in {directory}")

    # Process each file
    for file_path in docx_files:
        try:
            chunks = process_document(file_path, chunk_size, chunk_overlap)
            results[file_path.name] = chunks
        except Exception as e:
            # Log error but continue processing other files
            print(f"[WARNING] Failed to process {file_path.name}: {str(e)}")
            continue

    return results


def get_document_summary(chunks: List[Dict[str, str]]) -> Dict[str, any]:
    """
    Generate summary statistics for processed document chunks

    Args:
        chunks: List of processed chunks

    Returns:
        Dictionary with summary statistics
    """
    if not chunks:
        return {
            'total_chunks': 0,
            'total_characters': 0,
            'avg_chunk_size': 0
        }

    total_chars = sum(len(chunk['text']) for chunk in chunks)
    avg_chunk_size = total_chars / len(chunks) if chunks else 0

    return {
        'source': chunks[0]['source'] if chunks else 'unknown',
        'total_chunks': len(chunks),
        'total_characters': total_chars,
        'avg_chunk_size': int(avg_chunk_size),
        'avg_tokens': int(avg_chunk_size / 4)  # Rough estimate: 4 chars ≈ 1 token
    }
